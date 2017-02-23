from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from random import sample
from random import shuffle
from random import seed
 
nums = ['1','2','3','4','5','6','7','8','9','0']
met = [["Li",1],["Be",2],["Na",1],["Mg",2],["Al",3],["K",1],["Ca",2],["Ga",3]]
tmet = [["Ti",4],["Ti",3],["V",3],["V",5],["Cr",2],["Cr",3],["Mn",2],["Mn",4],["Fe",2],["Fe",3],["Co",3],["Co",4],["Ni",3],["Ni",2],["Cu",1],["Cu",2]]
nmet = [["N",3],["O",2],["F",1],["P",3],["S",2],["Cl",1],["Se",2],["Br",1],["I",1]]
pai = [["SO4",2],["SO3",2],["PO4",3],["PO3",3],["CO3",2],["NO3",1],["NO2",1]]
spec = [["Zn",2],["Ag",1]]
hyd = [["H",1]]
pref = ["Mono","Di","Tri","Tetra","Penta","Hexa","Hepta","Octa","Nona","Deca"]
rn = ["I","II","III","IV","V","VI","VII"]
 
enames = {"H": "Hydrogen",
          "He": "Helium",
          "Li": "Lithium",
          "Be": "Beryllium",
          "B": "Boron",
          "C": "Carbon",
          "N": "Nitrogen",
          "O": "Oxygen",
          "F": "Fluorine",
          "Ne": "Neon",
          "Na": "Sodium",
          "Mg": "Magnesium",
          "Al": "Aluminum",
          "Si": "Silicon",
          "P": "Phosphorus",
          "S": "Sulfur",
          "Cl": "Chlorine",
          "Ar": "Argon",
          "K": "Potassium",
          "Ca": "Calcium",
          "Sc": "Scandium",
          "Ti": "Titanium",
          "V": "Vanadium",
          "Cr": "Chromium",
          "Mn": "Manganese",
          "Fe": "Iron",
          "Co": "Cobalt",
          "Ni": "Nickel",
          "Cu": "Copper",
          "Zn": "Zinc",
          "Ga": "Gallium",
          "Ge": "Germanium",
          "As": "Arsenic",
          "Se": "Selenium",
          "Br": "Bromine",
          "Kr": "Krypton",
          "Rb": "Rubidium",
          "Sr": "Strontium",
          "Y": "Yttrium",
          "Zr": "Zirconium",
          "Nb": "Niobium",
          "Mo": "Molybdenum",
          "Tc": "Technetium",
          "Ru": "Ruthenium",
          "Rh": "Rhodium",
          "Pd": "Palladium",
          "Ag": "Silver",
          "Cd": "Cadmium",
          "In": "Indium",
          "Sn": "Tin",
          "Sb": "Antimony",
          "Te": "Tellurium",
          "I": "Iodine",
          "Xe": "Xenon",
          "Cs": "Cesium",
          "Ba": "Barium",
          "La": "Lanthanum",
          "Ce": "Cerium",
          "Pr": "Praseodymium",
          "Nd": "Neodymium",
          "Pm": "Promethium",
          "Sm": "Samarium",
          "Eu": "Europium",
          "Gd": "Gadolinium",
          "Tb": "Terbium",
          "Dy": "Dysprosium",
          "Ho": "Holmium",
          "Er": "Erbium",
          "Tm": "Thulium",
          "Yb": "Ytterbium",
          "Lu": "Lutetium",
          "Hf": "Hafnium",
          "Ta": "Tantalum",
          "W": "Tungsten",
          "Re": "Rhenium",
          "Os": "Osmium",
          "Ir": "Iridium",
          "Pt": "Platinum",
          "Au": "Gold",
          "Hg": "Mercury",
          "Tl": "Thallium",
          "Pb": "Lead",
          "Bi": "Bismuth",
          "Po": "Polonium",
          "At": "Astatine",
          "Rn": "Radon",
          "Fr": "Francium",
          "Ra": "Radium",
          "Ac": "Actinium",
          "Th": "Thorium",
          "Pa": "Protactinium",
          "U": "Uranium",
          "Np": "Neptunium",
          "Pu": "Plutonium",
          "Am": "Americium",
          "Cm": "Curium",
          "Bk": "Berkelium",
          "Cf": "Californium",
          "Es": "Einsteinium",
          "Fm": "Fermium",
          "Md": "Mendelevium",
          "No": "Nobelium",
          "Lr": "Lawrencium"}
 
sroots = {"C": "Carb",#standard roots
          "N": "Nitr",
          "O": "Ox",
          "F": "Fluor",
          "P": "Phosph",
          "S": "Sulf",
          "Cl": "Chlor",
          "Se": "Selen",
          "Br": "Brom",
          "I": "Iod"}
 
nmet_charge = {"C": 4,#standard roots
          "N": 3,
          "O": 2,
          "F": 1,
          "P": 3,
          "S": 2,
          "Cl": 1,
          "As": 3,
          "Se": 2,
          "Br": 1,
          "Te": 2,
          "I": 1}
 
painames = {"CO3": "Carbonate",#PAI roots
            "NO3": "Nitrate",
            "SO4": "Sulfate",
            "PO4": "Phosphate",
            "NO2": "Nitrite",
            "SO3": "Sulfite",
            "PO3": "Phosphite"}
 
aroots = {"C": "Carbon",#acid
          "CO3": "Carbon",
          "S": "Sulfur",
          "SO4": "Sulfur",
          "SO3": "Sulfur",
          "P": "Phosphor",
          "PO3": "Phosphor",
          "PO4": "Phosphor",
          "F": "Fluor",
          "Se": "Selen",
          "Br": "Brom",
          "I": "Iod",
          "Cl": "Chlor",
          "NO3": "Nitr",
          "NO2": "Nitr"}
 
        
type1A = []
type1B = []
type2A = []
type2B = []
type3 = []
type4A = []
type4B = []
typeSpecA = []
typeSpecB = []
 
type1A_names = []
type1B_names = []
type2A_names = []
type2B_names = []
type3_names = []
type4A_names = []
type4B_names = []
typeSpecA_names = []
typeSpecB_names = []
 
 
###Fill the type1A list
for a in met:
    for b in nmet:
        j = ""
        if max(a[1],b[1]) % min(a[1],b[1]) == 0:
            co1 = a[1]/min(a[1],b[1])
            co2 = b[1]/min(a[1],b[1])
            j = a[0]
            if co2 != 1:
                j += str(co2)
            j += b[0]
            if co1 != 1:
                j += str(co1)
        else:
            j = a[0] + str(b[1]) + b[0] + str(a[1])
        type1A.append(j)
 
###Fill the type1B list
for a in met:
    for b in pai:
        j = ""
        if max(a[1],b[1]) % min(a[1],b[1]) == 0:
            co1 = a[1]/min(a[1],b[1])
            co2 = b[1]/min(a[1],b[1])
            j = a[0]
            if co2 != 1:
                j += str(co2)
            if co1 != 1:
                j += "("
                j += b[0]
                j += ")"
                j += str(co1)
            else:
                j += b[0]
        else:
            j = a[0] + str(b[1]) + "(" + b[0] + ")" + str(a[1])
        type1B.append(j)
 
###Fill the type2A list
for a in tmet:
    for b in nmet:
        j = ""
        if max(a[1],b[1]) % min(a[1],b[1]) == 0:
            co1 = a[1]/min(a[1],b[1])
            co2 = b[1]/min(a[1],b[1])
            j = a[0]
            if co2 != 1:
                j += str(co2)
            j += b[0]
            if co1 != 1:
                j += str(co1)
        else:
            j = a[0] + str(b[1]) + b[0] + str(a[1])
        type2A.append(j)
 
###Fill the type2B list
for a in tmet:
    for b in pai:
        j = ""
        if max(a[1],b[1]) % min(a[1],b[1]) == 0:
            co1 = a[1]/min(a[1],b[1])
            co2 = b[1]/min(a[1],b[1])
            j = a[0]
            if co2 != 1:
                j += str(co2)
            if co1 != 1:
                j += "("
                j += b[0]
                j += ")"
                j += str(co1)
            else:
                j += b[0]
        else:
            j = a[0] + str(b[1]) + "(" + b[0] + ")" + str(a[1])
        type2B.append(j)
 
###Fill the type3A list
for a in nmet:
    for b in nmet:
        for c in range(1,11):
            for d in range(1,11):
                if a != b:
                    j = ""
                    j+=a[0]
                    if c > 1:
                        j += str(c)
                    j+=b[0]
                    if d > 1:
                        j += str(d)
                    if j != "AsS":
                        type3.append(j)
                
###Fill the type4A list
for a in hyd:
    for b in nmet:
        j = a[0]
        if b[1] > 1:
            j += str(b[1])
        j += b[0]
        if b[0] != "O":
            type4A.append(j)
 
###Fill the type4B list
for a in hyd:
    for b in pai:
        j = a[0]
        if b[1] > 1:
            j += str(b[1])
        j += b[0]
        type4B.append(j)
 
###Fill the SpecA list
for a in spec:
    for b in nmet:
        j = ""
        if max(a[1],b[1]) % min(a[1],b[1]) == 0:
            co1 = a[1]/min(a[1],b[1])
            co2 = b[1]/min(a[1],b[1])
            j = a[0]
            if co2 != 1:
                j += str(co2)
            j += b[0]
            if co1 != 1:
                j += str(co1)
        else:
            j = a[0] + str(b[1]) + b[0] + str(a[1])
        typeSpecA.append(j)
 
###Fill the SpecB list
for a in spec:
    for b in pai:
        j = ""
        if max(a[1],b[1]) % min(a[1],b[1]) == 0:
            co1 = a[1]/min(a[1],b[1])
            co2 = b[1]/min(a[1],b[1])
            j = a[0]
            if co2 != 1:
                j += str(co2)
            if co1 != 1:
                j += "("
                j += b[0]
                j += ")"
                j += str(co1)
            else:
                j += b[0]
        else:
            j = a[0] + str(b[1]) + "(" + b[0] + ")" + str(a[1])
        typeSpecB.append(j)
 
###Fill the names
for a in met:
    for b in nmet:
        type1A_names.append(enames[a[0]]+" "+sroots[b[0]] + "ide")
 
for a in met:
    for b in pai:
        type1B_names.append(enames[a[0]]+" "+painames[b[0]])
 
for a in tmet:
    for b in nmet:
        type2A_names.append(enames[a[0]]+" ("+rn[a[1]]+") "+sroots[b[0]] + "ide")
        
for a in tmet:
    for b in pai:
        type2B_names.append(enames[a[0]]+" ("+rn[a[1]]+") "+painames[b[0]])
 
for a in nmet:
    for b in nmet:
        for c in range(10):
            for d in range(10):
                if a != b and c != 0:
                    j = pref[c]+enames[a[0]].lower()+" "+pref[d]+sroots[b[0]].lower()+"ide"
                    j = j.replace("aa","a")
                    j = j.replace("ao","o")
                    j = j.replace("oo","o")
                    type3_names.append(j)
                if a != b and c == 0:
                    j = enames[a[0]]+" "+pref[d]+sroots[b[0]].lower()+"ide"
                    j = j.replace("aa","a")
                    j = j.replace("ao","o")
                    j = j.replace("oo","o")
                    type3_names.append(j)
 
for b in nmet[2:]:
    type4A_names.append("Hydro"+aroots[b[0]].lower()+"ic Acid")
 
for b in pai:
    if(painames[b[0]][-3:] == 'ate'):
        type4B_names.append(aroots[b[0]]+"ic Acid")
    else:
        type4B_names.append(aroots[b[0]]+"ous Acid")
 
for a in spec:
    for b in nmet:
        typeSpecA_names.append(enames[a[0]]+" "+sroots[b[0]] + "ide")
 
for a in spec:
    for b in pai:
        typeSpecB_names.append(enames[a[0]]+" "+painames[b[0]])
      
document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Tahoma'
font.size = Pt(10)
sections = document.sections
section = sections[0]
section.left_margin, section.right_margin = (Inches(0.5), Inches(0.5))
section.top_margin, section.bottom_margin = (Inches(0.5), Inches(0.5))
 
test_number = int(raw_input("Number of tests? "))
 
for version in range(test_number):
    document.add_paragraph('Name:____________________________________________________________________Date:________Hour:_______')
    document.add_paragraph('Unit 3 Chemistry Quiz #2 Version #'+str(version+1))
    document.add_paragraph('Write formulas for the following chemicals: ').bold = True
    names = []
    names.extend(sample(type1A_names,2))
    names.extend(sample(type1B_names,2))
    names.extend(sample(type2A_names,2))
    names.extend(sample(type2B_names,2))
    names.extend(sample(type3_names,2))
    shuffle(names)
    table = document.add_table(rows=5, cols=2)
    count = 0
    for row in table.rows:
        for cell in row.cells:
            cell.text = names[count]
            count += 1
    document.add_paragraph('Name the following chemicals: ').bold = True
    formulas = []
    formulas.extend(sample(type1A,2))
    formulas.extend(sample(type1B,2))
    formulas.extend(sample(type2A,2))
    formulas.extend(sample(type2B,2))
    formulas.extend(sample(type3,2))
    shuffle(formulas)
    table = document.add_table(rows=5, cols=2)
    for idx, name in enumerate(formulas):
        cells = table.rows[idx/2].cells
        paragraph = cells[idx%2].paragraphs[0]
        for x in name:
            if x in nums:
                run = paragraph.add_run(x)
                run.font.subscript = True
            else:
                run = paragraph.add_run(x)
    if version % 2 == 1:             
        document.add_page_break()
    else:
        document.add_paragraph('')
      
document.save('demo.doc')
