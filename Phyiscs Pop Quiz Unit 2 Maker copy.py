#!/usr/bin/env python
# -*- coding: UTF-8 -*-

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from random import sample
from random import shuffle
from random import seed
from random import randint

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)
    section.bottom_margin = Inches(.5)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)
tests = raw_input("How many tests? ")
for x in range(int(tests)):
    paragraph = document.add_paragraph('Name:_____________________________________________________________________________Date:_________________Hour:____________')
    table = document.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    a = table.cell(0,0)
    b = table.cell(0,1)
    A = a.merge(b)
    a = table.cell(2,0)
    b = table.cell(2,1)
    A = a.merge(b)
    a = table.cell(5,0)
    b = table.cell(5,1)
    A = a.merge(b)
    cell = table.cell(0,0)
    cell.text = "Teresita's car has broken down and is not moving!  She pushes the car on flat ground with a force of " + str(randint(400,600))+ "N for " +str(randint(101,299)/10.0) +" seconds. The car has a mass of "+ str(randint(3100,5900)) + " kilograms. Assume there is no friction." 
    cell = table.cell(1,0)
    run = cell.paragraphs[0].add_run("STEP 1: Calculate the value of F")
    run = cell.paragraphs[0].add_run('g')
    run.font.subscript = True
    run = cell.paragraphs[0].add_run('.  Show your work. Include the correct unit on your answer.\n\n\n\n')
    cell = table.cell(1,1)
    run = cell.paragraphs[0].add_run("STEP 2: What is the value of F")
    run = cell.paragraphs[0].add_run('N')
    run.font.subscript = True
    run = cell.paragraphs[0].add_run('?  Explain using at least two complete sentences.  Include the correct unit on your answer.\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
    cell = table.cell(2,0)
    cell.text = "STEP 3: Draw a Free Body Diagram for this scenario.  Assume there is no Friction.  Make sure your arrows are labeled and drawn at an appropriate length.\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n"
    cell = table.cell(3,0)
    run = cell.paragraphs[0].add_run("STEP 4: Calculate the value of F")
    run = cell.paragraphs[0].add_run('x')
    run.font.subscript = True
    run = cell.paragraphs[0].add_run('.  Show your work. Include the correct unit on your answer.\n\n\n\n\n\n\n')
    cell = table.cell(3,1)
    run = cell.paragraphs[0].add_run("STEP 5: Calculate the value of F")
    run = cell.paragraphs[0].add_run('y')
    run.font.subscript = True
    run = cell.paragraphs[0].add_run('.  Show your work. Include the correct unit on your answer.\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
    cell = table.cell(4,0)
    run = cell.paragraphs[0].add_run("STEP 6: Calculate the value of F")
    run = cell.paragraphs[0].add_run('NET')
    run.font.subscript = True
    run = cell.paragraphs[0].add_run('.  Show your work. Include the correct unit on your answer.\n\n\n\n\n\n\n')
    cell = table.cell(4,1)
    cell.text = "STEP 7: Calculate the acceleration of the car.  Show your work.  Include the correct unit in your answer.\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n"
    cell = table.cell(5,0)
    cell.text = "STEP 8: How far does Teresita push the car, in meters?  Show your work.  Include the correct unit in your answer.\n\n\n\n\n\n\n\n\n\n\n\n\n"
    
    document.add_page_break()



document.save('test.docx')
 
