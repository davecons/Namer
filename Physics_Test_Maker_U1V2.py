from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from random import sample
from random import shuffle
from random import seed
from random import randint

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)
    section.bottom_margin = Inches(.5)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)
tests = raw_input("How many tests? ")
for x in range(int(tests)):
    paragraph = document.add_paragraph('Name:_____________________________________________________________________________Date:_________________Hour:__________________')
    if x == 0:
        paragraph = document.add_paragraph('REVIEW SHEET UNIT 1 #2').bold = True
    else:
        paragraph = document.add_paragraph('UNIT 1 QUIZ #2 - Version ' + str(x)).bold = True
    table = document.add_table(rows=4, cols=2)
    a = table.cell(0,0)
    b = table.cell(0,1)
    A = a.merge(b)
    A.paragraphs[0].add_run("Alexis is riding her bike at " + str(randint(30,50)/10.0) + " m/s. She has a device on her bike that measures distance. When she notices she has traveled 150 meters, she slams on the brakes.  When she stops, she reads the device again - it reads " +str(randint(171,199)) + " meters. ")
    A.paragraphs[0].add_run("What was the acceleration of Alexis's bike?")
    cell = table.cell(1,0)
    cell.text = "STEP 1: Write the given values, with variables, for this problem.\n\n\n\n\n\n\n"
    cell = table.cell(1,1)
    cell.text = "STEP 2: Write the unknown variable"
    cell = table.cell(2,0)
    cell.text = "STEP 3: Write the equation you are going to use to answer this question.\n\n\n\n\n\n\n\n\n\n"
    cell = table.cell(2,1)
    cell.text = "STEP 4: Solve this equation for the variable you selected in STEP 2. Show your work."
    cell = table.cell(3,0)
    cell.text = "STEP 5: Evaluate this equation by plugging in your values from STEP 1.  Show your work.\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n"
    cell = table.cell(3,1)
    cell.text = "STEP 6: Write and circle your answer, with units."
    document.add_page_break()


    

document.save('test.docx')
 
