#!/usr/bin/env python
# -*- coding: UTF-8 -*-

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from random import sample
from random import shuffle
from random import seed
from random import randint

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)
    section.bottom_margin = Inches(.5)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)
tests = raw_input("How many tests? ")
for x in range(int(tests)):
    paragraph = document.add_paragraph('Name:_____________________________________________________________________________Date:_________________Hour:__________________')

    table = document.add_table(rows=4, cols=2)
    a = table.cell(0,0)
    b = table.cell(0,1)
    A = a.merge(b)
    cell = table.cell(0,0)
    cell.text = "Taya tries to throw a piece of paper into the trash.  She fires the paper at " + str(randint(80,100)/10.0)+ "m/s. The trash can is "+ str(randint(31,59)/10.0)+ " meters away.\nAt what angle must Taya throw the paper to hit the trash can?" 
    cell = table.cell(1,0)
    cell.text = "STEP 1: Write the given values, with variables, for this problem.\n\n\n\n\n\n\n"
    cell = table.cell(1,1)
    cell.text = "STEP 2: Write the unknown variable"
    cell = table.cell(2,0)
    cell.text = "STEP 3: Write the equation you are going to use to answer this question.\n\n\n\n\n\n\n\n\n\n"
    cell = table.cell(2,1)
    cell.text = "STEP 4: Solve this equation for the variable you selected in STEP 2. Show your work."
    cell = table.cell(3,0)
    cell.text = "STEP 5: Evaluate this equation by plugging in your values from STEP 1.  Show your work.\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n"
    cell = table.cell(3,1)
    cell.text = "STEP 6: Write and circle your answer, with units."
    document.add_page_break()



document.save('test.docx')
 
