from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from random import sample
from random import shuffle
from random import seed
from random import randint
from itertools import permutations as perm
from sympy.solvers import solve
from sympy import *
from sympy.solvers.solvers import check_assumptions
from sympy.parsing.sympy_parser import parse_expr


names = ["Ronnie","Teresita","Florine","Beaulah","Elizbeth","Ebony","Ronald","Luana","Zofia","Minerva","Eliz","Tegan","Ivelisse","Jeramy","Russel","Concha","Edna","Milan","Roma","Lester","Rachelle","Randi","Octavio","Denis","Anton","Loralee","Demarcus","Sharyl","Frieda","Evelynn","Angela","Lizabeth","Yuko","Alphonso","Carmelita","Aleshia","Roman","Melba","Delaine","Phylicia","Lashanda","Phylis","Jocelyn","Fransisca","Fabian","Luis","Libby","Donn","Joshua","Ozell"]
sports = ["running","skating","dancing","sliding","sledding","roller skating","skateboarding","swimming","diving","skiing","walking"]


t = Symbol('t')
vi = Symbol('vi')
vf = Symbol('vf')
a = Symbol('a')
di = Symbol('di')
df = Symbol('df')

eqs = [Eq(vf,vi+a*t),Eq(vf**2,vi**2+2*a*df),Eq(df,(vi+vf)/2*t)]
var = [t,vi,vf,a,df]
full_var = ['time','initial velocity','final velocity','acceleration','distance']
units = ['s','m/s','m/s','m/s^2','m']
for rounds in range(4):
    for eq in eqs:
        for v in var:
            if eq.has(v):
                seq = solve(eq,v)
                #print(eq,seq)
                j = str(seq[0]).split(',')
                #print(j[0])
                values = [0,0,0,0,0]
                if eq.has(t) and v != t:
                    values[0] = randint(1,10)/10.0
                if eq.has(vi) and v != vi:
                    values[1] = randint(10,30)/10.0
                if eq.has(vf) and v != vf:
                    values[2] = randint(70,90)/10.0
                if eq.has(a) and v != a:
                    values[3] = randint(10,30)/10.0
                if eq.has(df) and v != df:
                    values[4] = randint(10,30)/10.0
                if values[0] != 0:
                    print("t = ",values[0])
                if values[1] != 0:
                    print("vi = ",values[1])
                if values[2] != 0:
                    print("vf = ",values[2])
                if values[3] != 0:
                    print("a = ",values[3])
                if values[4] != 0:
                    print("df = ",values[4])
                print(str(v)+" = "+str(seq[0]),str(v)+" = "+str(abs(round(seq[0].subs({t:values[0],vi:values[1],vf:values[2],a:values[3],df:values[4]}),2))))
                pers = sample(names,1)
                print(str(pers[0])+" is "+str(sample(sports,1)[0])+".")
                '''pers = sample(names,1)
                question = ""
                question += str(pers)+" is "+str(sample(sports,1))+". "+str(pers)
                for x in range(5):
                    if values[x] != 0:
                        if full_var[x][0] == 'a' or full_var[x][0] == 'i':
                            question += " an "
                        else:
                            question += " a "
                    question += full_var[x]+" of "
                    question += str(values[x])
                    question += " "+str(units[x])
                    if x != 4:
                        question += ', '
                    print(question)'''
                        

ss = input("")




count = 0
total = 0
values = []
for group in problems:
    print(group)
    count += 1
    for o in perm(group,len(group)):
        y = list(o)
        total += 1
        temp = [0]*(len(y)-1)
        for x in range(len(temp)):
            if y[x] == 'initial velocity' or y[x] == 'initial distance':
                temp[x] = randint(10,30)/10.0
            else:
                temp[x] = randint(31,99)/10.0            
        print(temp,total,count,y[0:-1],y[-1])
        r = input("")

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)
    section.bottom_margin = Inches(.5)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)
tests = raw_input("How many tests? ")
for x in range(int(tests)):
    paragraph = document.add_paragraph('PHYSICS SEMESTER ONE EXAM\nDO NOT WRITE ON THIS PAPER')
    table = document.add_table(rows=20, cols=2)
    for y in range(15):
        for z in range(2):
            cell = table.cell(y,z)
            pick = randint(0,len(names)-1)
            student = names[pick]
            nums = []
            del names[pick]
            orig = sample(quants,5)
            use = sorted(sample(orig,4))
            rem = list(set(orig)-set(use))
            use.append(rem[0])
            question = cell.add_paragraph(student+" is "+str(''.join(sample(sports,1)))+". "+student+" has ")
            for a in range(4):
                if a == 3:
                    question.add_run("and ")
                if use[a][0] == 'a' or use[a][0] == 'i':
                    question.add_run("an ")
                else:
                    question.add_run("a ")
                nums.append(round(randint(10,99)/10.0,1))
                question.add_run(use[a] + " of "+str(nums[-1]))
                if use[a] == 'acceleration':
                    question.add_run(" m/s")
                    super_text = question.add_run('2')
                    super_text.font.superscript = True
                elif use[a] == 'initial velocity' or use[a] == 'final velocity':
                    question.add_run(" m/s")
                elif use[a] == 'initial distance' or use[a] == 'final distance':
                    question.add_run(" m")
                else:
                    question.add_run(" s")
                if a != 3:
                    question.add_run(", ")
            if use[4] == 'acceleration' or use[4] == 'initial velocity' or use[4] == 'final velocity':
                question.add_run(". What is "+student+"\'s "+use[4]+"?")
            elif use[4] == 'final distance':
                question.add_run(". How far does "+student+" travel, in meters?")
            else:
                question.add_run(". How long does "+student+" travel, in seconds?")
            J = sorted(use[0:4])
            K = use[4]
            print(J,K)
            print(nums)
            if J == ['acceleration', 'final distance', 'final velocity', 'initial velocity']:
                answer = 0

document.save('Physics Semester 1 Final Exam.docx')
 
